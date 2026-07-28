import hashlib
import io

import fitz
from docx import Document
from PIL import Image


class DocumentProcessor:
    # Critérios de filtragem de imagens (Fase 4 — pipeline de imagem+texto)
    MIN_LONG_SIDE_PX = 150
    MIN_SHORT_SIDE_PX = 80
    MAX_IMAGE_MB = 8
    MAX_IMAGES_PER_ANALYSIS = 12
    HEADER_FOOTER_ZONE_RATIO = 0.10  # 10% do topo/rodapé da página é ignorado

    @staticmethod
    def extract_plain_text(uploaded_file) -> str:
        ext = uploaded_file.name.split('.')[-1].lower()
        text = ""
        if ext == "pdf":
            data = uploaded_file.read()
            doc = fitz.open(stream=data, filetype="pdf")
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            uploaded_file.seek(0)  # deixa o stream pronto pra uma leitura seguinte (ex.: extração de imagens)
        elif ext == "docx":
            doc = Document(uploaded_file)
            for p in doc.paragraphs:
                text += p.text + "\n"
            uploaded_file.seek(0)
        elif ext == "txt":
            text = uploaded_file.getvalue().decode("utf-8")
        return text.strip()

    @staticmethod
    def extract_plain_text_multi(uploaded_files: list) -> str:
        """
        Extrai e concatena o texto de múltiplos arquivos, mantendo marcadores
        claros de onde cada documento começa/termina — isso ajuda a IA a não
        misturar contexto entre documentos diferentes (ex.: dois anexos que
        descrevem módulos distintos do mesmo sistema).
        """
        parts = []
        for uploaded_file in uploaded_files or []:
            text = DocumentProcessor.extract_plain_text(uploaded_file)
            if not text:
                continue
            parts.append(
                f"===== INÍCIO DO DOCUMENTO: {uploaded_file.name} =====\n"
                f"{text}\n"
                f"===== FIM DO DOCUMENTO: {uploaded_file.name} ====="
            )
        return "\n\n".join(parts).strip()

    # ------------------------------------------------------------------ #
    # Extração de imagens do CORPO do documento (Fase 4)
    #
    # Regras aplicadas, nessa ordem:
    #   1. Ignora imagens dentro da faixa de cabeçalho/rodapé (10% superior/
    #      inferior da página) — só PDF; em DOCX, cabeçalho/rodapé já são
    #      partes estruturalmente separadas do corpo, então nem entram aqui.
    #   2. Ignora imagens que se repetem em 2+ páginas/posições (logo/marca
    #      d'água colado manualmente, não um cabeçalho "de verdade").
    #   3. Ignora imagens menores que 150px no lado maior ou 80px no menor
    #      (ícones, elementos decorativos).
    #   4. Comprime imagens acima de 8MB (qualidade JPEG até um piso de 65,
    #      depois resolução até um piso de 1600px no lado maior); se mesmo
    #      assim continuar grande, descarta e registra aviso.
    #   5. Limita a 12 imagens por documento (na ordem em que aparecem).
    # ------------------------------------------------------------------ #
    @classmethod
    def extract_images_with_context(cls, uploaded_files: list) -> dict:
        """
        Retorna {"images": [...], "warnings": [...]}.
        Cada imagem: {
            "bytes": bytes (já comprimida se necessário),
            "mime": "image/jpeg" ou "image/png",
            "source_file": nome do arquivo de origem,
            "context": trecho de texto próximo à imagem (ajuda a IA a entender o contexto),
            "location": descrição legível (ex.: "página 3"),
        }
        """
        all_images = []
        warnings = []

        for uploaded_file in uploaded_files or []:
            ext = uploaded_file.name.split('.')[-1].lower()
            if ext == "pdf":
                imgs, warns = cls._extract_images_pdf(uploaded_file)
            elif ext == "docx":
                imgs, warns = cls._extract_images_docx(uploaded_file)
            else:
                imgs, warns = [], []
            all_images.extend(imgs)
            warnings.extend(warns)

        if len(all_images) > cls.MAX_IMAGES_PER_ANALYSIS:
            warnings.append(
                f"O documento tinha {len(all_images)} imagens elegíveis; só as primeiras "
                f"{cls.MAX_IMAGES_PER_ANALYSIS} foram consideradas (limite por análise)."
            )
            all_images = all_images[:cls.MAX_IMAGES_PER_ANALYSIS]

        return {"images": all_images, "warnings": warnings}

    @classmethod
    def _extract_images_pdf(cls, uploaded_file):
        warnings = []
        data = uploaded_file.read()
        uploaded_file.seek(0)
        doc = fitz.open(stream=data, filetype="pdf")

        raw_candidates = []
        hash_pages = {}

        for page_num, page in enumerate(doc):
            page_h = page.rect.height
            header_limit = page_h * cls.HEADER_FOOTER_ZONE_RATIO
            footer_limit = page_h * (1 - cls.HEADER_FOOTER_ZONE_RATIO)
            page_text = page.get_text()

            for img in page.get_images(full=True):
                xref = img[0]
                try:
                    rects = page.get_image_rects(xref)
                    if not rects:
                        continue
                    rect = rects[0]
                    # inteiramente dentro da faixa de cabeçalho ou rodapé -> ignora
                    if rect.y1 <= header_limit or rect.y0 >= footer_limit:
                        continue

                    base_image = doc.extract_image(xref)
                    img_bytes = base_image["image"]
                    width, height = base_image.get("width", 0), base_image.get("height", 0)
                    long_side, short_side = max(width, height), min(width, height)
                    if long_side < cls.MIN_LONG_SIDE_PX or short_side < cls.MIN_SHORT_SIDE_PX:
                        continue

                    img_hash = hashlib.md5(img_bytes).hexdigest()
                    hash_pages.setdefault(img_hash, set()).add(page_num)

                    raw_candidates.append({
                        "hash": img_hash,
                        "bytes": img_bytes,
                        "ext": base_image.get("ext", "png"),
                        "page": page_num + 1,
                        "context": page_text[:600],
                        "source_file": uploaded_file.name,
                    })
                except Exception as error:
                    warnings.append(f"{uploaded_file.name}: falha ao processar uma imagem (pulada): {error}")

        doc.close()

        # remove imagens repetidas em 2+ páginas (provável logo/marca d'água solta no corpo)
        repeated = {h for h, pages in hash_pages.items() if len(pages) >= 2}
        filtered = [c for c in raw_candidates if c["hash"] not in repeated]

        images = []
        for c in filtered:
            compressed, mime, warn = cls._compress_if_needed(c["bytes"], c["ext"])
            if warn:
                warnings.append(f"{uploaded_file.name} (página {c['page']}): {warn}")
            if compressed is None:
                continue
            images.append({
                "bytes": compressed,
                "mime": mime,
                "source_file": c["source_file"],
                "context": c["context"],
                "location": f"página {c['page']}",
            })
        return images, warnings

    @classmethod
    def _extract_images_docx(cls, uploaded_file):
        """
        Extração pro DOCX: cabeçalho/rodapé já ficam de fora automaticamente
        (são partes estruturalmente separadas do documento no formato OOXML,
        nunca aparecem em document.paragraphs/document.part principal).
        """
        warnings = []
        doc = Document(uploaded_file)
        uploaded_file.seek(0)

        hash_count = {}
        raw_candidates = []

        # Mapa aproximado de contexto: concatena o texto dos parágrafos
        # próximos de onde a imagem foi encontrada na relação do documento.
        full_text = "\n".join(p.text for p in doc.paragraphs)

        for rel_id, rel in doc.part.rels.items():
            if "image" not in rel.reltype:
                continue
            try:
                image_part = rel.target_part
                img_bytes = image_part.blob
                with Image.open(io.BytesIO(img_bytes)) as pil_img:
                    width, height = pil_img.size
                long_side, short_side = max(width, height), min(width, height)
                if long_side < cls.MIN_LONG_SIDE_PX or short_side < cls.MIN_SHORT_SIDE_PX:
                    continue

                img_hash = hashlib.md5(img_bytes).hexdigest()
                hash_count[img_hash] = hash_count.get(img_hash, 0) + 1

                ext = image_part.content_type.split("/")[-1]
                raw_candidates.append({
                    "hash": img_hash,
                    "bytes": img_bytes,
                    "ext": ext,
                    "source_file": uploaded_file.name,
                })
            except Exception as error:
                warnings.append(f"{uploaded_file.name}: falha ao processar uma imagem embutida (pulada): {error}")

        repeated = {h for h, count in hash_count.items() if count >= 2}
        filtered = [c for c in raw_candidates if c["hash"] not in repeated]

        images = []
        for c in filtered:
            compressed, mime, warn = cls._compress_if_needed(c["bytes"], c["ext"])
            if warn:
                warnings.append(f"{uploaded_file.name}: {warn}")
            if compressed is None:
                continue
            images.append({
                "bytes": compressed,
                "mime": mime,
                "source_file": c["source_file"],
                "context": full_text[:600],  # aproximação: contexto geral do documento
                "location": "corpo do documento",
            })
        return images, warnings

    @classmethod
    def _compress_if_needed(cls, img_bytes: bytes, ext: str):
        """
        Retorna (bytes_finais, mime_type, aviso|None). Se a imagem já está
        dentro do limite, devolve como está. Se precisar comprimir, tenta
        reduzir qualidade JPEG (piso 65) e depois resolução (piso 1600px no
        lado maior) antes de desistir e descartar.
        """
        max_bytes = cls.MAX_IMAGE_MB * 1024 * 1024
        if len(img_bytes) <= max_bytes:
            mime = "image/png" if ext.lower() == "png" else "image/jpeg"
            return img_bytes, mime, None

        try:
            with Image.open(io.BytesIO(img_bytes)) as pil_img:
                pil_img = pil_img.convert("RGB")

                for quality in (90, 80, 70, 65):
                    buf = io.BytesIO()
                    pil_img.save(buf, format="JPEG", quality=quality, optimize=True)
                    if buf.tell() <= max_bytes:
                        return buf.getvalue(), "image/jpeg", None

                width, height = pil_img.size
                long_side = max(width, height)
                min_long_side = 1600
                while long_side > min_long_side:
                    scale = 0.85
                    new_w, new_h = int(width * scale), int(height * scale)
                    if max(new_w, new_h) < min_long_side:
                        new_scale = min_long_side / max(width, height)
                        new_w, new_h = int(width * new_scale), int(height * new_scale)
                    pil_img = pil_img.resize((new_w, new_h), Image.LANCZOS)
                    width, height = new_w, new_h
                    long_side = max(width, height)

                    buf = io.BytesIO()
                    pil_img.save(buf, format="JPEG", quality=65, optimize=True)
                    if buf.tell() <= max_bytes:
                        return buf.getvalue(), "image/jpeg", None

                return None, None, (
                    f"Imagem descartada — mesmo após compressão, continuou acima de "
                    f"{cls.MAX_IMAGE_MB}MB sem ficar ilegível."
                )
        except Exception as error:
            return None, None, f"Falha ao comprimir imagem, descartada: {error}"
